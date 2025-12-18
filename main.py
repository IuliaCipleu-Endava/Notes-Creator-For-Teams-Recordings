import sys
import nltk
import docx
from pathlib import Path
from collections import Counter

nltk.download('punkt', force=False)

# ------------------------------
# CONFIGURE YOUR FOLDERS HERE
# ------------------------------

TRANSCRIPTS_FOLDER = Path(r"C:\Users\icipleu\OneDrive - ENDAVA\Transcripts")
ONENOTE_FOLDER = Path(r"C:\Users\icipleu\OneDrive - ENDAVA\Documents\OneNote Notebooks\Meetings")

# ------------------------------
# SUPPORT FUNCTIONS
# ------------------------------

def extract_keywords(text, num_keywords=5):
    words = nltk.word_tokenize(text)
    words = [w.lower() for w in words if w.isalnum() and len(w) > 3]
    freq = Counter(words)
    return [w for w, _ in freq.most_common(num_keywords)]

def abstractive_summary(text):
    try:
        from transformers import pipeline
        summarizer = pipeline('summarization', model='facebook/bart-large-cnn')
        result = summarizer(text, max_length=130, min_length=30, do_sample=False)
        return result[0]['summary_text']
    except Exception:
        return None

def read_latest_docx(folder: Path):
    files = list(folder.glob("*.docx"))
    if not files:
        print("No .docx transcript found.")
        sys.exit(1)

    newest = max(files, key=lambda f: f.stat().st_mtime)
    doc = docx.Document(newest)
    content = "\n".join(p.text for p in doc.paragraphs).strip()
    return newest.name, content

# ------------------------------
# MAIN PROCESS
# ------------------------------

# 1. Read latest transcript
filename, input_text = read_latest_docx(TRANSCRIPTS_FOLDER)
print(f"Processing transcript: {filename}")
if not input_text:
    print("Transcript is empty.")
    sys.exit(0)

# 2. Extractive summary
sentences = nltk.sent_tokenize(input_text)
extractive = " ".join(sentences[:5])
print("Extractive summary created.")
# 3. Try abstractive summary
# abstractive = abstractive_summary(input_text)
# print("Abstractive summary created." if abstractive else "Abstractive summary not available.")
# 4. Keywords
keywords = extract_keywords(input_text)
print(f"Keywords extracted: {keywords}")
# 5. Build summary text
output = []
output.append("--- Notes Summary ---")

if extractive:
    output.append("\nAbstractive Summary:\n" + extractive)
else:
    output.append("\nExtractive Summary:\n" + extractive)

output.append("\nKey Points:")
for kw in keywords:
    output.append(f"- {kw}")

final_text = "\n".join(output)


# 6. Ensure OneNote folder exists
ONENOTE_FOLDER.mkdir(parents=True, exist_ok=True)

# 7. Save to OneNote folder as .txt
save_path_txt = ONENOTE_FOLDER / f"Summary-{filename.replace('.docx','')}.txt"
save_path_txt.write_text(final_text, encoding="utf-8")

# 8. Also save as .docx
save_path_docx = ONENOTE_FOLDER / f"Summary-{filename.replace('.docx','')}.docx"
doc = docx.Document()
for line in output:
    doc.add_paragraph(line)
doc.save(save_path_docx)

print(f"Summary created: {save_path_txt}")
print(f"Summary also created as DOCX: {save_path_docx}")
print("OneNote will sync and auto-create the page.")
