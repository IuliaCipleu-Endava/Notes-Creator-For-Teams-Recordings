import sys
import json
import nltk
import docx
import re
from pathlib import Path
from collections import Counter
from langdetect import detect, DetectorFactory
import dateparser
from datetime import datetime
from http.server import BaseHTTPRequestHandler, HTTPServer

from llama_cpp import Llama  # GGUF / llama.cpp backend

DetectorFactory.seed = 0
nltk.download("punkt", quiet=False)

# -------------------------------------------------------------------
# 1. MODEL LOADING (GGUF / llama.cpp)
# -------------------------------------------------------------------
# Change this to your actual GGUF path
GGUF_MODEL_PATH = r"C:\Users\icipleu\OneDrive - ENDAVA\Documents\Data Project\Notes_Creator_For_Teams_Recordings\dolphin-2_6-phi-2_oasst2_chatml_v2.q4_k_m.gguf"

print("Loading GGUF model (phi-2)…")
llm = Llama(
    model_path=GGUF_MODEL_PATH,
    n_ctx=2048,  # match training context
    n_threads=6,
    n_gpu_layers=0,
    verbose=False,
)

print("Model loaded.")

# -------------------------------------------------------------------
# 2. CONFIG / WORDLISTS
# -------------------------------------------------------------------
CONFIG_DIR = Path("config")


def load_wordlist(path: Path):
    if not path.exists():
        return []
    with open(path, encoding="utf-8") as f:
        return [ln.strip() for ln in f if ln.strip()]


STOPWORDS_RO = set(load_wordlist(CONFIG_DIR / "stopwords_ro.txt"))
STOPWORDS_EN = set(load_wordlist(CONFIG_DIR / "stopwords_en.txt"))

TODO_RO = load_wordlist(CONFIG_DIR / "todo_keywords_ro.txt")
TODO_EN = load_wordlist(CONFIG_DIR / "todo_keywords_en.txt")

SOLVED_RO = load_wordlist(CONFIG_DIR / "solved_keywords_ro.txt")
SOLVED_EN = load_wordlist(CONFIG_DIR / "solved_keywords_en.txt")

SUGGEST_RO = load_wordlist(CONFIG_DIR / "suggestions_keywords_ro.txt")
SUGGEST_EN = load_wordlist(CONFIG_DIR / "suggestions_keywords_en.txt")

ISSUES_RO = load_wordlist(CONFIG_DIR / "issues_keywords_ro.txt")
ISSUES_EN = load_wordlist(CONFIG_DIR / "issues_keywords_en.txt")

# -------------------------------------------------------------------
# 3. HELPERS
# -------------------------------------------------------------------
def detect_language(text: str) -> str:
    try:
        return "romanian" if detect(text) == "ro" else "english"
    except Exception:
        return "english"

def count_tokens(text):
    return len(llm.tokenize(text.encode("utf-8")))

def chunk_text_llm_safe(text, max_tokens=1400):
    """
    Split text into chunks guaranteed to fit the model's context window.
    """
    sentences = nltk.sent_tokenize(text)
    chunks = []
    current = ""

    for s in sentences:
        cand = (current + " " + s).strip()
        tok = len(llm.tokenize(cand.encode("utf-8")))

        if tok > max_tokens:
            if current:
                chunks.append(current)
            current = s
        else:
            current = cand

    if current:
        chunks.append(current)
    print(f"Text chunked into {len(chunks)} parts for LLM processing.")
    return chunks



def extract_keywords(text: str, language="english", n=5):
    words = nltk.word_tokenize(text.lower())
    words = [w for w in words if w.isalnum()]
    stop = STOPWORDS_RO if language == "romanian" else STOPWORDS_EN
    words = [w for w in words if w not in stop and len(w) > 3]
    freq = Counter(words)
    return [w for w, _ in freq.most_common(n)]


def clean_transcript(text: str) -> str:
    lines = text.split("\n")
    lines = [ln.strip() for ln in lines if ln.strip()]
    lines = [ln for ln in lines if "transcription" not in ln.lower()]
    cleaned = []
    for ln in lines:
        ln = re.sub(r"\s+\d+:\d+", "", ln).strip()
        cleaned.append(ln)
    return "\n".join(cleaned)


def categorize_sentences(text: str, language: str):
    todos, solved, suggestions, issues = [], [], [], []
    sentences = nltk.sent_tokenize(text)

    if language == "romanian":
        todo_words = TODO_RO
        solved_words = SOLVED_RO
        suggestion_words = SUGGEST_RO
        issue_words = ISSUES_RO
    else:
        todo_words = TODO_EN
        solved_words = SOLVED_EN
        suggestion_words = SUGGEST_EN
        issue_words = ISSUES_EN

    for s in sentences:
        low = s.lower()
        if any(w in low for w in todo_words):
            todos.append(s)
            continue
        if any(w in low for w in solved_words):
            solved.append(s)
            continue
        if any(w in low for w in suggestion_words):
            suggestions.append(s)
            continue
        if any(w in low for w in issue_words):
            issues.append(s)
            continue

    return todos, solved, suggestions, issues


def extract_dates(text: str, language: str):
    """
    Return list of (sentence, date_str) for important-ish dates found.
    """
    results = []
    sentences = nltk.sent_tokenize(text)
    lang_code = "ro" if language == "romanian" else "en"

    for s in sentences:
        dt = dateparser.parse(
            s,
            settings={"PREFER_DATES_FROM": "future"},
            languages=[lang_code],
        )
        if dt:
            results.append((s, dt.strftime("%Y-%m-%d %H:%M")))
    return results


def read_vtt(path: Path) -> str:
    """
    Clean Teams .vtt files:
    - Remove headers, NOTE, cue IDs, timecodes
    - Extract <v Speaker>Text</v> to "Speaker: Text"
    """
    lines = []
    current_speaker = None

    with open(path, encoding="utf-8") as f:
        for raw in f:
            line = raw.strip()

            if not line:
                continue
            if line == "WEBVTT":
                continue
            if line.startswith("NOTE"):
                continue
            if line.isdigit():  # numeric cue IDs
                continue
            if re.match(r"\d{2}:\d{2}:\d{2}\.\d{3} -->", line):
                continue

            m = re.match(r"<v ([^>]+)>(.*?)</v>", line)
            if m:
                speaker = m.group(1).strip().split(".")[0]
                text = m.group(2).strip()
                lines.append(f"{speaker}: {text}")
                current_speaker = speaker
                continue

            m = re.match(r"<v ([^>]+)>(.*)", line)
            if m:
                current_speaker = m.group(1).strip().split(".")[0]
                rest = m.group(2).strip()
                if rest:
                    lines.append(f"{current_speaker}: {rest}")
                continue

            if current_speaker:
                lines.append(f"{current_speaker}: {line}")
                continue

            lines.append(line)

    cleaned = [ln.replace("</v>", "").strip() for ln in lines if ln.strip()]
    return "\n".join(cleaned)


def read_docx(path: Path) -> str:
    doc = docx.Document(path)
    raw = "\n".join(p.text.strip() for p in doc.paragraphs if p.text.strip())
    raw = re.sub(r"started transcription", "", raw, flags=re.I)
    raw = re.sub(r"stopped transcription", "", raw, flags=re.I)
    raw = re.sub(r"\b\d{1,2}:\d{2}\b", "", raw)
    raw = re.sub(r"\n{2,}", "\n", raw)
    return raw.strip()


def extract_participants(text: str):
    participants = set()
    for line in text.splitlines():
        m = re.match(r"^([A-ZȘȚĂÂÎ][\wăâîșț]+(?: [A-ZȘȚĂÂÎ][\wăâîșț]+)*)\s*:", line)
        if m:
            participants.add(m.group(1))
            continue
    return sorted(participants)

def normalize_json_structure(parsed):
    """
    Accepts both:
    {
      "summary": "...",
      "todo": [...]
    }

    and:
    {
      "meeting_notes": {
          "summary": "...",
          "todo": [...]
      }
    }

    Returns flat dict:
    {
      "summary": "...",
      "todo": [...],
      ...
    }
    """
    if "meeting_notes" in parsed and isinstance(parsed["meeting_notes"], dict):
        # flatten
        return parsed["meeting_notes"]

    return parsed  # already flat

def extract_first_json(text):
    start = text.find("{")
    if start == -1:
        return None

    stack = 0
    for i in range(start, len(text)):
        if text[i] == "{":
            stack += 1
        elif text[i] == "}":
            stack -= 1
            if stack == 0:
                try:
                    return json.loads(text[start:i+1])
                except:
                    return None
    return None

def extract_structured_from_any_json(obj, result):
    """
    Recursively merge structured meeting fields from ANY JSON shape.
    Supports:
    - nested dicts
    - lists of dicts
    - meeting_notes wrappers
    - title + lists
    - JSON with missing keys
    """
    if isinstance(obj, dict):
        for key in ["summary", "todo", "solved", "issues", "suggestions", "decisions"]:
            if key in obj:
                val = obj[key]
                if key == "summary" and isinstance(val, str):
                    result["summary"].append(val)
                elif isinstance(val, list):
                    result[key].extend(v for v in val if isinstance(v, str))
        for v in obj.values():
            extract_structured_from_any_json(v, result)

    elif isinstance(obj, list):
        for item in obj:
            extract_structured_from_any_json(item, result)

def extract_from_explanation(expl, key):
    """
    Extract sections like:
    Summary: ...
    Todo: ...
    Issues: ...
    """
    pattern = rf"{key}:(.*?)(?=\n[A-Z][a-z]+:|$)"
    m = re.search(pattern, expl, re.S)
    if not m:
        return []
    lines = [ln.strip(" -") for ln in m.group(1).split("\n") if ln.strip()]
    # remove garbage lines like "Rest"
    return [l for l in lines if l.lower() not in ("rest",)]

def llama_structured_notes(clean_text: str, language: str, raw_output_path: str) -> dict:
    """
    Extract structured meeting notes using llama.cpp in ChatML mode.
    Ensures valid JSON output and merges across multiple chunks.
    """

    MAX_CHUNK_TOKENS = 1400
    chunks = chunk_text_llm_safe(clean_text, max_tokens=MAX_CHUNK_TOKENS)

    combined = {k: [] for k in
                ["summary", "todo", "solved", "issues", "suggestions", "decisions"]}

    # --- Utility functions --------------------------------------------------

    def strip_fences(text: str) -> str:
        """Remove ```json or ``` wrappers."""
        text = text.strip()
        text = re.sub(r"^```(?:json)?\s*", "", text, flags=re.I)
        text = re.sub(r"```$", "", text)
        return text.strip()

    def safe_json_extract(text: str):
        """Return first valid JSON found, or None."""
        text = strip_fences(text)
        return extract_first_json(text)

    # --- Chat prompts -------------------------------------------------------

    SYSTEM_PROMPT = (
        "You are an AI assistant that extracts structured JSON from meeting transcripts. "
        "You MUST return ONLY valid JSON. "
        "No explanations, no commentary, no text outside of JSON."
    )

    JSON_SCHEMA = """
{
  "meeting_notes": {
    "summary": "",
    "todo": [],
    "solved": [],
    "issues": [],
    "suggestions": [],
    "decisions": []
  }
}
"""

    # --- Main loop ----------------------------------------------------------

    for idx, chunk in enumerate(chunks):

        USER_PROMPT = f"""
LANGUAGE: {language}

Extract structured meeting notes from the following transcript chunk.

Return ONLY valid JSON in this exact structure:
{JSON_SCHEMA}

MEETING TRANSCRIPT:
\"\"\"{chunk}\"\"\"
"""

        # Call llama in ChatML mode
        res = llm.create_chat_completion(
            messages=[
                {"role": "system", "content": SYSTEM_PROMPT},
                {"role": "user", "content": USER_PROMPT}
            ],
            max_tokens=400,
            temperature=0.0,
        )

        raw = res["choices"][0]["message"]["content"]
        raw_clean = strip_fences(raw)

        print(f"\n=== RAW OUTPUT CHUNK {idx+1} ===")
        print(raw)

        if raw_output_path:
            with open(raw_output_path, "a", encoding="utf-8") as f:
                f.write(f"\n=== CHUNK {idx+1} ===\n{raw}\n")

        # Attempt JSON extraction
        parsed = safe_json_extract(raw_clean)
        if parsed:
            extract_structured_from_any_json(parsed, combined)
            print(f"Chunk {idx+1}: JSON extracted.")
        else:
            print(f"Chunk {idx+1}: No JSON extracted.")

    # --- Final merge and cleanup -------------------------------------------

    final = {
        "summary": " ".join(dict.fromkeys([s for s in combined["summary"] if s.strip()])),
        "todo": list(dict.fromkeys([t for t in combined["todo"] if t.strip()])),
        "solved": list(dict.fromkeys([s for s in combined["solved"] if s.strip()])),
        "issues": list(dict.fromkeys([i for i in combined["issues"] if i.strip()])),
        "suggestions": list(dict.fromkeys([g for g in combined["suggestions"] if g.strip()])),
        "decisions": list(dict.fromkeys([d for d in combined["decisions"] if d.strip()])),
    }

    return final


# -------------------------------------------------------------------
# 4. CORE PROCESSING (what used to be main())
# -------------------------------------------------------------------
def process_meetings(input_folder: str, output_folder: str,
                     previous_count: int, current_count: int):

    in_dir = Path(input_folder)
    out_dir = Path(output_folder)
    out_dir.mkdir(parents=True, exist_ok=True)

    n_files = max(current_count - previous_count, 1)

    # latest N .docx + .vtt
    files = list(in_dir.glob("*.docx")) + list(in_dir.glob("*.vtt"))
    files.sort(key=lambda f: f.stat().st_mtime, reverse=True)
    files_to_process = files[:n_files]

    processed = []

    for f in files_to_process:
        print(f"Processing: {f.name}")

        if f.suffix.lower() == ".docx":
            raw = read_docx(f)
        else:
            raw = read_vtt(f)

        if not raw.strip():
            print("Empty transcript.")
            continue

        clean = clean_transcript(raw)
        language = detect_language(clean)
        participants = extract_participants(clean)
        keywords = extract_keywords(clean, language)
        dates = extract_dates(clean, language)
        todos_h, solved_h, sugg_h, issues_h = categorize_sentences(clean, language)

        # Ask GGUF model for structured info
        raw_path = out_dir / f"Summary-{f.stem}.raw.txt"
        struct = llama_structured_notes(clean, language, raw_output_path=raw_path)

        # Merge heuristic + model (model preferred, heuristics as backup)
        summary_list = struct.get("summary", [])
        if isinstance(summary_list, list):
            summary = " ".join(summary_list).strip()
        else:
            summary = summary_list.strip() or " ".join(nltk.sent_tokenize(clean)[:5])

        todos = struct.get("todo", []) or todos_h
        solved = struct.get("solved", []) or solved_h
        issues = struct.get("issues", []) or issues_h
        suggestions = struct.get("suggestions", []) or sugg_h
        decisions = struct.get("decisions", [])

        # Build text output
        txt_lines = []
        txt_lines.append("--- Meeting Notes Summary ---")
        txt_lines.append(f"Language: {language.capitalize()}")
        txt_lines.append(
            f"Participants: {', '.join(participants) if participants else 'N/A'}"
        )
        txt_lines.append("")
        txt_lines.append("Key Points:")
        txt_lines.extend(f"- {kw}" for kw in keywords)
        txt_lines.append("")
        txt_lines.append("Summary:")
        txt_lines.append(summary)
        txt_lines.append("")
        txt_lines.append("To Do:")
        txt_lines.extend(f"- {t}" for t in todos) if todos else txt_lines.append(
            "- None identified"
        )
        txt_lines.append("")
        txt_lines.append("Solved / Completed:")
        txt_lines.extend(f"- {s}" for s in solved) if solved else txt_lines.append(
            "- None mentioned"
        )
        txt_lines.append("")
        txt_lines.append("Suggestions:")
        txt_lines.extend(f"- {s}" for s in suggestions) if suggestions else txt_lines.append(
            "- None identified"
        )
        txt_lines.append("")
        txt_lines.append("Issues:")
        txt_lines.extend(f"- {i}" for i in issues) if issues else txt_lines.append(
            "- No issues flagged"
        )
        txt_lines.append("")
        txt_lines.append("Decisions:")
        txt_lines.extend(f"- {d}" for d in decisions) if decisions else txt_lines.append(
            "- None identified"
        )
        txt_lines.append("")
        txt_lines.append("Important Dates:")
        if dates:
            for sent, d in dates:
                txt_lines.append(f"- {d}: {sent}")
        else:
            txt_lines.append("- No important dates mentioned")

        txt_content = "\n".join(txt_lines)
        base = f"Summary-{f.stem}"

        # Save TXT
        txt_path = out_dir / f"{base}.txt"
        txt_path.write_text(txt_content, encoding="utf-8")

        # Save DOCX
        doc_out = docx.Document()
        doc_out.add_heading("Meeting Notes Summary", level=0)
        doc_out.add_heading("Language", level=1)
        doc_out.add_paragraph(language.capitalize())
        doc_out.add_heading("Participants", level=1)
        doc_out.add_paragraph(", ".join(participants) if participants else "N/A")

        doc_out.add_heading("Key Points", level=1)
        for kw in keywords:
            doc_out.add_paragraph(kw, style="List Bullet")

        doc_out.add_heading("Summary", level=1)
        doc_out.add_paragraph(summary)

        doc_out.add_heading("To Do", level=1)
        if todos:
            for t in todos:
                doc_out.add_paragraph(t, style="List Bullet")
        else:
            doc_out.add_paragraph("No action items identified.")

        doc_out.add_heading("Solved / Completed", level=1)
        if solved:
            for s in solved:
                doc_out.add_paragraph(s, style="List Bullet")
        else:
            doc_out.add_paragraph("None mentioned.")

        doc_out.add_heading("Suggestions", level=1)
        if suggestions:
            for s in suggestions:
                doc_out.add_paragraph(s, style="List Bullet")
        else:
            doc_out.add_paragraph("No suggestions mentioned.")

        doc_out.add_heading("Issues", level=1)
        if issues:
            for i in issues:
                doc_out.add_paragraph(i, style="List Bullet")
        else:
            doc_out.add_paragraph("No issues detected.")

        doc_out.add_heading("Decisions", level=1)
        if decisions:
            for d in decisions:
                doc_out.add_paragraph(d, style="List Bullet")
        else:
            doc_out.add_paragraph("No decisions mentioned.")

        doc_out.add_heading("Important Dates", level=1)
        if dates:
            for sent, d in dates:
                doc_out.add_paragraph(f"{d}: {sent}", style="List Bullet")
        else:
            doc_out.add_paragraph("No important dates mentioned.")

        docx_path = out_dir / f"{base}.docx"
        doc_out.save(docx_path)

        processed.append(
            {
                "file": str(f),
                "summary_txt": str(txt_path),
                "summary_docx": str(docx_path),
            }
        )

    return processed

# -------------------------------------------------------------------
# 5. SIMPLE HTTP SERVER FOR PAD
# -------------------------------------------------------------------
class Handler(BaseHTTPRequestHandler):
    def _send_json(self, obj, status=200):
        data = json.dumps(obj).encode("utf-8")
        self.send_response(status)
        self.send_header("Content-Type", "application/json; charset=utf-8")
        self.send_header("Content-Length", str(len(data)))
        self.end_headers()
        self.wfile.write(data)

    def do_POST(self):
        if self.path != "/process":
            self._send_json({"error": "Not found"}, status=404)
            return

        length = int(self.headers.get("Content-Length", 0))
        body = self.rfile.read(length)
        try:
            payload = json.loads(body.decode("utf-8"))
        except Exception:
            self._send_json({"error": "Invalid JSON"}, status=400)
            return

        try:
            input_folder = payload["input_folder"]
            output_folder = payload["output_folder"]
            prev = int(payload.get("previous_count", 0))
            curr = int(payload.get("current_count", 0))
        except KeyError as e:
            self._send_json({"error": f"Missing field {e}"}, status=400)
            return

        try:
            processed = process_meetings(input_folder, output_folder, prev, curr)
            self._send_json({"status": "ok", "processed": processed})
        except Exception as e:
            self._send_json({"status": "error", "message": str(e)}, status=500)


def run_server(port=8000):
    server = HTTPServer(("127.0.0.1", port), Handler)
    print(f"Server listening on http://127.0.0.1:{port}/process")
    server.serve_forever()


if __name__ == "__main__":
    run_server()
