import json
import nltk
import docx
import re
from pathlib import Path
from collections import Counter
from langdetect import detect
import dateparser

from llama_cpp import Llama  # GGUF / llama.cpp backend
from functools import lru_cache
print("Loading Qwen tokenizer and model…")
from transformers import AutoModelForCausalLM, AutoTokenizer
import torch

@lru_cache(maxsize=10000)
def cached_tokenize(text: str):
    """Tokenize text with LLM tokenizer using caching."""
    return tuple(llm.tokenize(text.encode("utf-8")))


# -------------------------------------------------------------------
# 1. MODEL LOADING (GGUF / llama.cpp)
# -------------------------------------------------------------------
# Change this to your actual GGUF path
GGUF_MODEL_PATH = r"C:\Users\icipleu\OneDrive - ENDAVA\Documents\Data Project\Notes_Creator_For_Teams_Recordings\dolphin-2_6-phi-2_oasst2_chatml_v2.q4_k_m.gguf"

print("Loading GGUF model (phi-2)…")
llm = Llama(
    model_path=GGUF_MODEL_PATH,
    n_ctx=2048,
    n_threads=6,
    n_gpu_layers=0,
    chat_format="chatml",
    verbose=False,
)
print("Model loaded.")

model_name = "Qwen/Qwen2.5-7B-Instruct"

tokenizer = AutoTokenizer.from_pretrained(model_name, trust_remote_code=True)
model = AutoModelForCausalLM.from_pretrained(
    model_name,
    trust_remote_code=True,
    device_map="auto",
    dtype=torch.float16,   # important
)
print("Model loaded.")

# -------------------------------------------------------------------
# 2. CONFIG / WORDLISTS
# -------------------------------------------------------------------
CONFIG_DIR = Path("config")


def load_wordlist(path: Path):
    if not path.exists():
        return []
    with open(path, encoding="utf-8") as f:
        return [ln.strip() for ln in f if ln.strip()]


STOPWORDS_RO = set(load_wordlist(CONFIG_DIR / "stopwords_ro.txt"))
STOPWORDS_EN = set(load_wordlist(CONFIG_DIR / "stopwords_en.txt"))

TODO_RO = load_wordlist(CONFIG_DIR / "todo_keywords_ro.txt")
TODO_EN = load_wordlist(CONFIG_DIR / "todo_keywords_en.txt")

SOLVED_RO = load_wordlist(CONFIG_DIR / "solved_keywords_ro.txt")
SOLVED_EN = load_wordlist(CONFIG_DIR / "solved_keywords_en.txt")

SUGGEST_RO = load_wordlist(CONFIG_DIR / "suggestions_keywords_ro.txt")
SUGGEST_EN = load_wordlist(CONFIG_DIR / "suggestions_keywords_en.txt")

ISSUES_RO = load_wordlist(CONFIG_DIR / "issues_keywords_ro.txt")
ISSUES_EN = load_wordlist(CONFIG_DIR / "issues_keywords_en.txt")

TODO_PHRASES = [
    r"\bwe need to\b",
    r"\bshould\b",
    r"\bmust\b",
    r"\bto do\b",
    r"\baction item\b",
    r"\blet's\b",
    r"\bde facut\b",
]

ISSUE_PHRASES = [
    r"\bproblem\b",
    r"\berror\b",
    r"\bblocked\b",
    r"\bchallenge\b",
    r"\bissue\b",
    r"\bnu merge\b",
]

SUGGESTION_PHRASES = [
    r"\bpropose\b",
    r"\bsuggest\b",
    r"\brecommend\b",
    r"\bar fi bine\b",
]

FILLER_WORDS_REGEX = r"\b(yes|ok|mhm|uh|hm|right|yeah|alright|a bit)\b\.?"
ADMIN_CHATTER_REGEX = r"(can you hear me|let me share|hold on|wait a second|starting recording|stopped recording).*"

def detect_language(text: str) -> str:
    try:
        return "romanian" if detect(text) == "ro" else "english"
    except Exception:
        return "english"

def count_tokens(text):
    return len(llm.tokenize(text.encode("utf-8")))

@lru_cache(maxsize=5000)
def cached_sent_tokenize(text: str):
    return tuple(nltk.sent_tokenize(text))


# def chunk_text_llm_safe(text, max_tokens=1400):
#     """
#     Split text into chunks guaranteed to fit the model's context window.
#     """
#     sentences = cached_sent_tokenize(text)
#     chunks = []
#     current = ""

#     for s in sentences:
#         cand = (current + " " + s).strip()
#         tok = len(cached_tokenize(cand))

#         if tok > max_tokens:
#             if current:
#                 chunks.append(current)
#             current = s
#         else:
#             current = cand

#     if current:
#         chunks.append(current)
#     print(f"Text chunked into {len(chunks)} parts for LLM processing.")
#     return chunks
def chunk_text_llm_safe(text, max_sentences=25):
    sentences = cached_sent_tokenize(text)
    chunks = []

    for i in range(0, len(sentences), max_sentences):
        chunk = " ".join(sentences[i:i + max_sentences])
        if chunk.strip():
            chunks.append(chunk)

    print(f"Text chunked into {len(chunks)} parts.")
    return chunks


def extract_keywords(text: str, language="english", n=5):
    words = nltk.word_tokenize(text.lower())
    words = [w for w in words if w.isalnum()]
    stop = STOPWORDS_RO if language == "romanian" else STOPWORDS_EN
    words = [w for w in words if w not in stop and len(w) > 3]
    freq = Counter(words)
    return [w for w, _ in freq.most_common(n)]


def clean_transcript(text: str) -> str:
    lines = text.split("\n")
    lines = [ln.strip() for ln in lines if ln.strip()]
    lines = [ln for ln in lines if "transcription" not in ln.lower()]
    cleaned = []
    for ln in lines:
        ln = re.sub(r"\s+\d+:\d+", "", ln).strip()
        cleaned.append(ln)
    return "\n".join(cleaned)


def categorize_sentences(text: str, language: str):
    todos, solved, suggestions, issues = [], [], [], []
    sentences = cached_sent_tokenize(text)

    if language == "romanian":
        todo_words = TODO_RO
        solved_words = SOLVED_RO
        suggestion_words = SUGGEST_RO
        issue_words = ISSUES_RO
    else:
        todo_words = TODO_EN
        solved_words = SOLVED_EN
        suggestion_words = SUGGEST_EN
        issue_words = ISSUES_EN

    for s in sentences:
        low = s.lower()
        if any(w in low for w in todo_words):
            todos.append(s)
            continue
        if any(w in low for w in solved_words):
            solved.append(s)
            continue
        if any(w in low for w in suggestion_words):
            suggestions.append(s)
            continue
        if any(w in low for w in issue_words):
            issues.append(s)
            continue

    return todos, solved, suggestions, issues


def extract_dates(text: str, language: str):
    """
    Return list of (sentence, date_str) for important-ish dates found.
    """
    results = []
    sentences = cached_sent_tokenize(text)
    lang_code = "ro" if language == "romanian" else "en"

    for s in sentences:
        dt = dateparser.parse(
            s,
            settings={"PREFER_DATES_FROM": "future"},
            languages=[lang_code],
        )
        if dt:
            results.append((s, dt.strftime("%Y-%m-%d %H:%M")))
    return results


def read_vtt(path: Path) -> str:
    """
    Robust cleaner for Microsoft Teams .vtt transcripts.
    Handles:
    - WEBVTT headers, NOTE blocks
    - Cue IDs, numeric lines
    - Timecodes
    - <v Speaker>…</v> blocks
    - Continuation lines for speakers
    - Removes GUID fragments, formatting artifacts, <c> tags
    - Removes zero-width spaces & stray tokens
    - Ignores garbage-only lines
    """
    lines = []
    current_speaker = None

    guid_pattern = re.compile(
        r"[0-9a-f]{8}-[0-9a-f]{4}-[0-9a-f]{4}-[0-9a-f]{4}-[0-9a-f]{12}(?:/\d+-\d+)?",
        re.I
    )

    with open(path, encoding="utf-8") as f:
        for raw in f:
            line = raw.strip()

            # Skip empties, headers, NOTES
            if not line or line == "WEBVTT" or line.startswith("NOTE"):
                continue

            # Remove zero-width spaces
            line = line.replace("\u200b", "")

            # Remove speaker styling <c.colorXXXXXX>
            line = re.sub(r"<c[^>]*>", "", line)
            line = line.replace("</c>", "")

            # Remove GUID fragments
            line = guid_pattern.sub("", line).strip()

            # Skip pure garbage lines like: 11-0, 0/0, 39-1
            if re.fullmatch(r"[\d/.-]+", line):
                continue

            # Skip timecodes (00:00:00.000 --> 00:00:00.000)
            if re.match(r"\d{2}:\d{2}:\d{2}\.\d{3} -->", line):
                continue

            # Remove inline timestamps inside text
            line = re.sub(r"<\d{2}:\d{2}:\d{2}\.\d{3}>", "", line)

            # Full <v>…</v> on one line
            m = re.match(r"<v ([^>]+)>(.*?)</v>", line)
            if m:
                speaker = m.group(1).split(".")[0].strip()
                text = m.group(2).strip()
                if text:
                    lines.append(f"{speaker}: {text}")
                current_speaker = speaker
                continue

            # Opening <v Speaker> without closing
            m = re.match(r"<v ([^>]+)>(.*)", line)
            if m:
                speaker = m.group(1).split(".")[0].strip()
                text = m.group(2).strip()
                current_speaker = speaker
                if text:
                    lines.append(f"{speaker}: {text}")
                continue

            # Continuation line for current speaker
            if current_speaker:
                if line not in ("</v>",):
                    lines.append(f"{current_speaker}: {line}")
                continue

            # Fallback: treat as generic text line
            lines.append(line)

    # Cleanup: remove empty lines, stray </v>, double prefixes
    cleaned = []
    for ln in lines:
        ln = ln.replace("</v>", "").strip()

        # Remove accidental duplication: "Speaker: Speaker: text"
        ln = re.sub(r"^([^:]+):\s*\1:\s*", r"\1: ", ln)

        if ln:
            cleaned.append(ln)

    return "\n".join(cleaned)

def read_docx(path: Path) -> str:
    doc = docx.Document(path)
    raw = "\n".join(p.text.strip() for p in doc.paragraphs if p.text.strip())
    raw = re.sub(r"started transcription", "", raw, flags=re.I)
    raw = re.sub(r"stopped transcription", "", raw, flags=re.I)
    raw = re.sub(r"\b\d{1,2}:\d{2}\b", "", raw)
    raw = re.sub(r"\n{2,}", "\n", raw)
    return raw.strip()

def extract_participants(text: str):
    """
    Extract participant names from cleaned DOCX or VTT transcripts.
    DOCX format example:
        Name   0:06
    VTT format example:
        Name: text
    """
    participants = set()
    lines = text.splitlines()

    # ------------------------------------------------------
    # MODE DETECTION
    # ------------------------------------------------------
    # DOCX transcripts almost always contain timestamps like "0:06", "12:44"
    has_docx_timestamps = any(re.search(r"\b\d{1,2}:\d{2}\b", ln) for ln in lines)

    # VTT transcripts contain "Name:" patterns
    has_vtt_colon = any(re.match(r"^[A-ZȘȚĂÂÎ][\wăâîșț]+.*:", ln) for ln in lines)

    # ------------------------------------------------------
    # REGEXES
    # ------------------------------------------------------

    # DOCX speaker format:
    #   First Last   0:45
    pattern_docx = re.compile(
        r"^([A-ZȘȚĂÂÎ][a-zA-ZăâîșțȘȚĂÂÎ]+(?: [A-ZȘȚĂÂÎ][a-zA-ZăâîșțȘȚĂÂÎ]+){0,2})\s+\d{1,2}:\d{2}\b"
    )

    # VTT speaker format:
    #   First Last: text
    pattern_vtt = re.compile(
        r"^([A-ZȘȚĂÂÎ][\wăâîșț]+(?: [A-ZȘȚĂÂÎ][\wăâîșț]+){0,2})\s*:",
    )

    # ------------------------------------------------------
    # EXTRACTION LOGIC
    # ------------------------------------------------------
    if has_docx_timestamps:
        # Extract DOCX-style speakers
        for ln in lines:
            m = pattern_docx.match(ln)
            if m:
                participants.add(m.group(1).strip())

    if has_vtt_colon or not participants:
        # Extract VTT-style speakers
        for ln in lines:
            m = pattern_vtt.match(ln)
            if m:
                participants.add(m.group(1).strip())

    # Cleanup: remove duplicates, sort
    return sorted(participants)


def normalize_json_structure(parsed):
    """
    Accepts both:
    {
      "summary": "...",
      "todo": [...]
    }

    and:
    {
      "meeting_notes": {
          "summary": "...",
          "todo": [...]
      }
    }

    Returns flat dict:
    {
      "summary": "...",
      "todo": [...],
      ...
    }
    """
    if "meeting_notes" in parsed and isinstance(parsed["meeting_notes"], dict):
        # flatten
        return parsed["meeting_notes"]

    return parsed  # already flat

def extract_first_json(text):
    start = text.find("{")
    if start == -1:
        return None

    stack = 0
    for i in range(start, len(text)):
        if text[i] == "{":
            stack += 1
        elif text[i] == "}":
            stack -= 1
            if stack == 0:
                try:
                    return json.loads(text[start:i+1])
                except:
                    return None
    return None

def extract_structured_from_any_json(obj, result):
    """
    Recursively merge structured meeting fields from ANY JSON shape.
    Supports:
    - nested dicts
    - lists of dicts
    - meeting_notes wrappers
    - title + lists
    - JSON with missing keys
    """
    if isinstance(obj, dict):
        for key in ["summary", "todo", "solved", "issues", "suggestions", "decisions"]:
            if key in obj:
                val = obj[key]
                if key == "summary" and isinstance(val, str):
                    result["summary"].append(val)
                elif isinstance(val, list):
                    result[key].extend(v for v in val if isinstance(v, str))
        for v in obj.values():
            extract_structured_from_any_json(v, result)

    elif isinstance(obj, list):
        for item in obj:
            extract_structured_from_any_json(item, result)

def extract_from_explanation(expl, key):
    """
    Extract sections like:
    Summary: ...
    Todo: ...
    Issues: ...
    """
    pattern = rf"{key}:(.*?)(?=\n[A-Z][a-z]+:|$)"
    m = re.search(pattern, expl, re.S)
    if not m:
        return []
    lines = [ln.strip(" -") for ln in m.group(1).split("\n") if ln.strip()]
    # remove garbage lines like "Rest"
    return [l for l in lines if l.lower() not in ("rest",)]

def filter_items(items):
    cleaned = []
    for item in items:
        if not isinstance(item, str):
            continue
        low = item.lower()
        if "meeting transcript" in low or "meeting notes" in low:
            continue
        cleaned.append(item.strip())
    return cleaned

def process_meetings(input_folder: str, output_folder: str,
                     previous_count: int, current_count: int):

    in_dir = Path(input_folder)
    out_dir = Path(output_folder)
    out_dir.mkdir(parents=True, exist_ok=True)

    n_files = max(current_count - previous_count, 1)

    # latest N .docx + .vtt
    files = list(in_dir.glob("*.docx")) + list(in_dir.glob("*.vtt"))
    files.sort(key=lambda f: f.stat().st_mtime, reverse=True)
    files_to_process = files[:n_files]

    processed = []

    for f in files_to_process:
        print(f"Processing: {f.name}")

        if f.suffix.lower() == ".docx":
            raw = read_docx(f)
        else:
            raw = read_vtt(f)

        if not raw.strip():
            print("Empty transcript.")
            continue

        clean = clean_transcript(raw)
        language = detect_language(clean)
        participants = extract_participants(clean)
        participants = resolve_aliases(participants, clean)
        keywords = extract_keywords(clean, language)
        dates = extract_dates(clean, language)
        todos_h, solved_h, sugg_h, issues_h = categorize_sentences(clean, language)

        # Ask GGUF model for structured info
        raw_path = out_dir / f"Summary-{f.stem}.raw.txt"
        struct = llama_structured_notes_resilient(clean, language, raw_output_path=raw_path)

        # Merge heuristic + model (model preferred, heuristics as backup)
        summary_list = struct.get("summary", [])
    
        if isinstance(summary_list, list):
            summary_list = filter_items(summary_list)
            summary = merge_summaries(summary_list)
        else:
            summary = summary_list.strip() or " ".join(cached_sent_tokenize(clean)[:5])
            if "meeting transcript" in summary.lower() or "meeting notes" in summary.lower():
                summary = ""
        summary = clean_summary(summary)

        # Remove participant names from summary if any participants found
        if participants and summary:
            for name in participants:
                summary = re.sub(rf"(?:^|(?<=\.|\n))\s*{re.escape(name)}:\s*", "", summary)

        keywords = struct.get("keywords", []) or keywords
        todos_h2, issues_h2, sugg_h2 = heuristic_extract_items(clean)
        todos = filter_items(struct.get("todo", []) or todos_h or todos_h2)
        issues = filter_items(struct.get("issues", []) or issues_h or issues_h2)
        suggestions = filter_items(struct.get("suggestions", []) or sugg_h or sugg_h2)
        solved = filter_items(struct.get("solved", []) or solved_h)
        decisions = filter_items(struct.get("decisions", []))
        deadlines = extract_deadlines(clean)
        stats = speaker_stats(clean)
        top_speaker = stats.most_common(1)[0][0] if stats else "N/A"
        ai_suggestion = struct.get("ai_suggestion", "").strip()

        # Build text output
        txt_lines = []
        txt_lines.append("--- Meeting Notes Summary ---")
        txt_lines.append(f"Language: {language.capitalize()}")
        txt_lines.append(
            f"Participants: {', '.join(participants) if participants else 'N/A'}"
        )
        txt_lines.append("")
        txt_lines.append("Key Points:")
        txt_lines.extend(f"- {kw}" for kw in keywords)
        txt_lines.append("")
        txt_lines.append("Summary:")
        txt_lines.append(summary)
        txt_lines.append("")
        txt_lines.append("To Do:")
        txt_lines.extend(f"- {t}" for t in todos) if todos else txt_lines.append(
            "- None identified"
        )
        txt_lines.append("")
        txt_lines.append("Solved / Completed:")
        txt_lines.extend(f"- {s}" for s in solved) if solved else txt_lines.append(
            "- None mentioned"
        )
        txt_lines.append("")
        txt_lines.append("Suggestions:")
        txt_lines.extend(f"- {s}" for s in suggestions) if suggestions else txt_lines.append(
            "- None identified"
        )
        txt_lines.append("")
        txt_lines.append("Issues:")
        txt_lines.extend(f"- {i}" for i in issues) if issues else txt_lines.append(
            "- No issues flagged"
        )
        txt_lines.append("")
        txt_lines.append("Decisions:")
        txt_lines.extend(f"- {d}" for d in decisions) if decisions else txt_lines.append(
            "- None identified"
        )
        txt_lines.append("")
        txt_lines.append("Important Dates:")
        if dates:
            for sent, d in dates:
                txt_lines.append(f"- {d}: {sent}")
        else:
            txt_lines.append("- No important dates mentioned")
        txt_lines.append("Deadlines:")
        if deadlines:
            txt_lines.extend(f"- {d}" for d in deadlines)
        else:
            txt_lines.append("- None detected")
        txt_lines.append("")
        txt_lines.append(f"Most active speaker: {top_speaker}")
        if ai_suggestion:
            txt_lines.append("")
            txt_lines.append("AI Suggestion:")
            txt_lines.append(ai_suggestion)
        
        txt_content = "\n".join(txt_lines)
        base = f"Summary-{f.stem}"

        # Save TXT
        txt_path = out_dir / f"{base}.txt"
        txt_path.write_text(txt_content, encoding="utf-8")

        # Save DOCX
        doc_out = docx.Document()
        doc_out.add_heading("Meeting Notes Summary", level=0)
        doc_out.add_heading("Language", level=1)
        doc_out.add_paragraph(language.capitalize())
        doc_out.add_heading("Participants", level=1)
        doc_out.add_paragraph(", ".join(participants) if participants else "N/A")

        doc_out.add_heading("Key Points", level=1)
        for kw in keywords:
            doc_out.add_paragraph(kw, style="List Bullet")

        doc_out.add_heading("Summary", level=1)
        doc_out.add_paragraph(summary)

        doc_out.add_heading("To Do", level=1)
        if todos:
            for t in todos:
                doc_out.add_paragraph(t, style="List Bullet")
        else:
            doc_out.add_paragraph("No action items identified.")

        doc_out.add_heading("Solved / Completed", level=1)
        if solved:
            for s in solved:
                doc_out.add_paragraph(s, style="List Bullet")
        else:
            doc_out.add_paragraph("None mentioned.")

        doc_out.add_heading("Suggestions", level=1)
        if suggestions:
            for s in suggestions:
                doc_out.add_paragraph(s, style="List Bullet")
        else:
            doc_out.add_paragraph("No suggestions mentioned.")

        doc_out.add_heading("Issues", level=1)
        if issues:
            for i in issues:
                doc_out.add_paragraph(i, style="List Bullet")
        else:
            doc_out.add_paragraph("No issues detected.")

        doc_out.add_heading("Decisions", level=1)
        if decisions:
            for d in decisions:
                doc_out.add_paragraph(d, style="List Bullet")
        else:
            doc_out.add_paragraph("No decisions mentioned.")


        doc_out.add_heading("Important Dates", level=1)
        if dates:
            for sent, d in dates:
                doc_out.add_paragraph(f"{d}: {sent}", style="List Bullet")
        else:
            doc_out.add_paragraph("No important dates mentioned.")

        doc_out.add_heading("Deadlines", level=1)
        if deadlines:
            for d in deadlines:
                doc_out.add_paragraph(d, style="List Bullet")
        else:
            doc_out.add_paragraph("None detected.")

        doc_out.add_heading("Most active speaker", level=1)
        doc_out.add_paragraph(top_speaker)
        
        if ai_suggestion:
            doc_out.add_heading("AI Suggestion", level=1)
            doc_out.add_paragraph(ai_suggestion)

        docx_path = out_dir / f"{base}.docx"
        doc_out.save(docx_path)

        processed.append(
            {
                "file": str(f),
                "summary_txt": str(txt_path),
                "summary_docx": str(docx_path),
            }
        )

    return processed

# ==========================================
#  ENHANCED PARTICIPANT HANDLING
# ==========================================

def extract_participants(text: str):
    """
    Extract participant names from DOCX-style timestamps and VTT-style "Name:" lines.
    Handles:
    - First Last   0:06
    - First Middle Last   12:45
    - First Last: text
    - Romanian diacritics
    """

    participants = set()
    lines = text.splitlines()

    # Detect docx vs vtt mode
    has_timestamp = any(re.search(r"\b\d{1,2}:\d{2}\b", ln) for ln in lines)
    has_colon = any(re.match(r"^[A-ZȘȚĂÂÎ][\wăâîșț]+.*:", ln) for ln in lines)

    # DOCX pattern: "Firstname Lastname   0:45"
    pattern_docx = re.compile(
        r"^([A-ZȘȚĂÂÎ][a-zA-ZăâîșțȘȚĂÂÎ]+(?: [A-ZȘȚĂÂÎ][a-zA-ZăâîșțȘȚĂÂÎ]+){0,2})\s+\d{1,2}:\d{2}\b"
    )

    # VTT pattern: "Firstname Lastname: ..."
    pattern_vtt = re.compile(
        r"^([A-ZȘȚĂÂÎ][\wăâîșț]+(?: [A-ZȘȚĂÂÎ][\wăâîșț]+){0,2})\s*:"
    )

    # DOCX extraction
    if has_timestamp:
        for ln in lines:
            m = pattern_docx.match(ln)
            if m:
                participants.add(m.group(1).strip())

    # VTT extraction
    if has_colon or not participants:
        for ln in lines:
            m = pattern_vtt.match(ln)
            if m:
                participants.add(m.group(1).strip())

    return sorted(participants)


# ==========================================
#  PARTICIPANT ALIAS RESOLUTION
# ==========================================

def resolve_aliases(participants, text):
    """
    Convert short names into full names.
    """
    if not participants:
        return participants

    first_name_map = {p.split()[0]: p for p in participants}
    resolved = set()

    for ln in text.splitlines():
        token = ln.split(":")[0].strip()
        if token in first_name_map:
            resolved.add(first_name_map[token])

    return sorted(resolved or participants)


# ==========================================
#  SUMMARY CLEANUP MODULE
# ==========================================

def clean_summary(summary):
    """
    Remove filler words, meeting-control dialog,
    repeated acknowledgments, and noise.
    """

    if not summary:
        return summary

    # Remove conversational fillers
    summary = re.sub(
        FILLER_WORDS_REGEX,
        "",
        summary,
        flags=re.I,
    )

    # Remove meeting admin chatter
    summary = re.sub(
        FILLER_WORDS_REGEX,
        "",
        summary,
        flags=re.I
    )

    # Remove leftover speaker prefixes
    summary = re.sub(r"^[A-ZȘȚĂÂÎ][\wăâîșț]+:", "", summary)

    # Remove duplicate whitespace
    summary = " ".join(summary.split())

    return summary.strip()


# ==========================================
#  IMPROVED HEURISTIC TASK + ISSUE FINDER
# ==========================================

def heuristic_extract_items(text):
    todos, issues, suggestions = [], [], []
    sentences = cached_sent_tokenize(text)

    for s in sentences:
        low = s.lower()
        todo_phrases = TODO_EN + TODO_RO
        issue_phrases = ISSUES_EN + ISSUES_RO
        suggestion_phrases = SUGGEST_EN + SUGGEST_RO
        if any(re.search(p, low) for p in todo_phrases):
            todos.append(s)

        if any(re.search(p, low) for p in issue_phrases):
            issues.append(s)

        if any(re.search(p, low) for p in suggestion_phrases):
            suggestions.append(s)

    return todos, issues, suggestions


# ==========================================
#  DEADLINE & DATE IMPROVEMENTS
# ==========================================

DEADLINE_PATTERNS = [
    r"\bby (monday|tuesday|wednesday|thursday|friday|tomorrow|end of month)\b",
    r"\bbefore \d{1,2}/\d{1,2}\b",
    r"\bdue\b",
    r"\buntil\b",
]

def extract_deadlines(text):
    deadlines = []
    for pattern in DEADLINE_PATTERNS:
        for m in re.finditer(pattern, text, flags=re.I):
            deadlines.append(m.group())
    return deadlines


# ==========================================
#  SPEAKER STATISTICS
# ==========================================

def speaker_stats(text):
    """
    Count how many speaking turns each participant had.
    Useful for identifying who dominates the meeting.
    """
    stats = Counter()
    for ln in text.splitlines():
        if ":" in ln:
            speaker = ln.split(":")[0].strip()
            stats[speaker] += 1
    return stats


# ==========================================
#  SUMMARY MERGING IMPROVEMENT
# ==========================================

def merge_summaries(summary_list):
    """
    Remove redundant sentences and keep only meaningful content.
    """
    seen = set()
    final = []

    for s in summary_list:
        s = s.strip()
        if not s:
            continue
        if s.lower() in seen:
            continue

        seen.add(s.lower())
        final.append(s)

    return " ".join(final)

###########Will test Monday####################
import math
from collections import Counter

def _word_set(text, n_min_len=3):
    words = re.findall(r"\w+", text.lower())
    return set(w for w in words if len(w) >= n_min_len)

def _overlap_ratio(small_text, large_text):
    """
    proportion of words in small_text that also appear in large_text
    (used to verify model output actually derives from chunk)
    """
    sset = _word_set(small_text)
    lset = _word_set(large_text)
    if not sset:
        return 0.0
    return len(sset & lset) / len(sset)

def _truncate_after_json(text):
    """
    Truncate text after the first complete top-level JSON object (if any),
    otherwise return original.
    """
    start = text.find("{")
    if start == -1:
        return text
    stack = 0
    for i in range(start, len(text)):
        if text[i] == "{":
            stack += 1
        elif text[i] == "}":
            stack -= 1
            if stack == 0:
                return text[:i+1]
    return text

def strip_speaker_chatter(text: str) -> str:
    lines = []
    for ln in text.splitlines():
        ln = ln.strip()

        # remove "Name:" prefixes
        ln = re.sub(r"^[A-Z][a-z]+(?: [A-Z][a-z]+)*:\s*", "", ln)

        # drop pure chatter
        if ln.lower() in {
            "ok", "okay", "yeah", "yes", "no",
            "can you see the screen",
            "can you hear me",
            "well", "so", "uh", "um"
        }:
            continue

        # drop very short noise
        if len(ln.split()) < 4:
            continue

        lines.append(ln)

    return "\n".join(lines)

def extract_all_json_objects(text: str):
    """
    Extract ALL JSON objects from text, even if surrounded by garbage.
    Returns a list of parsed dicts.
    """
    results = []
    stack = []
    start = None

    for i, ch in enumerate(text):
        if ch == "{":
            if not stack:
                start = i
            stack.append("{")
        elif ch == "}":
            if stack:
                stack.pop()
                if not stack and start is not None:
                    candidate = text[start:i+1]
                    try:
                        results.append(json.loads(candidate))
                    except Exception:
                        pass
                    start = None

    return results

def merge_structured_content(json_objs, combined):
    """
    Merge any recognizable meeting fields from ANY JSON shape.
    """
    FIELDS = {
        "summary": str,
        "bullets": list,
        "keywords": list,
        "todo": list,
        "issues": list,
        "suggestions": list,
        "decisions": list,
        "solved": list,
    }

    def walk(obj):
        if isinstance(obj, dict):
            for k, v in obj.items():
                if k in FIELDS:
                    if isinstance(v, FIELDS[k]):
                        if k == "summary":
                            combined["summary"].append(v)
                        else:
                            combined.setdefault(k, []).extend(
                                x for x in v if isinstance(x, str)
                            )
                walk(v)
        elif isinstance(obj, list):
            for it in obj:
                walk(it)

    for o in json_objs:
        walk(o)

def salvage_from_text(raw, combined):
    for line in raw.splitlines():
        line = line.strip()
        low = line.lower()

        if len(line.split()) > 12 and not any(x in low for x in [
            "can you", "did you", "are you",
            "screen", "call", "access",
            "calendar", "hello"
        ]):
            combined["summary"].append(line)

        if line.startswith(("-", "*")):
            combined["bullets"].append(line.lstrip("-* ").strip())


def llama_structured_notes_resilient(clean_text: str, language: str, raw_output_path: str) -> dict:
    """
    Robust structured-note extraction for noisy meeting transcripts.

    Key design changes vs the previous version:
    - No per-chunk "overlap validation" that discards useful content.
    - Two-phase extraction:
        (1) Chunk-level: extract factual bullets + candidate items (todo/issues/suggestions/decisions) in JSON.
        (2) Final pass: synthesize ONE clean structured JSON from all collected bullets/items.
    - Heuristics are always used as a baseline (seed), then LLM refines.
    - Aggressive garbage filtering (e.g., "Meeting transcript summary") is removed or softened.
    """

    MAX_CHUNK_TOKENS = 1400
    clean_text = strip_speaker_chatter(clean_text)

    chunks = chunk_text_llm_safe(clean_text, max_tokens=MAX_CHUNK_TOKENS)

    # -----------------------------
    # Helpers
    # -----------------------------
    def _append_raw(label: str, text: str):
        if not raw_output_path:
            return
        with open(raw_output_path, "a", encoding="utf-8") as f:
            f.write(f"\n=== {label} ===\n{text}\n")

    def _call_llm(system_prompt: str, user_prompt: str, max_tokens: int = 500) -> str:
        try:
            if hasattr(llm, "create_chat_completion"):
                res = llm.create_chat_completion(
                    messages=[
                        {"role": "system", "content": system_prompt},
                        {"role": "user", "content": user_prompt},
                    ],
                    max_tokens=max_tokens,
                    temperature=0.0,
                    stop=["\n\nNote:", "\n\nNOTES:", "\n\nExplanation:", "Note:"],
                )
                return res["choices"][0]["message"]["content"]
            # fallback non-chat
            res = llm(user_prompt, max_tokens=max_tokens, temperature=0.0)
            if isinstance(res, dict) and "choices" in res and res["choices"] and "text" in res["choices"][0]:
                return res["choices"][0]["text"]
            return str(res)
        except Exception as e:
            return f"LLM ERROR: {e}"
        
    def _call_qwen(system_prompt: str, user_prompt: str, max_tokens: int = 500) -> str:
        prompt = system_prompt + "\n" + user_prompt

        inputs = tokenizer(
            prompt,
            return_tensors="pt",
            truncation=True
        )

        inputs = {k: v.to(model.device) for k, v in inputs.items()}

        with torch.no_grad():
            outputs = model.generate(
                **inputs,
                max_new_tokens=max_tokens,
                do_sample=False,
                temperature=0.0,
                eos_token_id=tokenizer.eos_token_id,
            )

        return tokenizer.decode(outputs[0], skip_special_tokens=True)


    def _strip_fences(text: str) -> str:
        text = (text or "").strip()
        text = re.sub(r"^```(?:json)?\s*", "", text, flags=re.I)
        text = re.sub(r"\s*```$", "", text)
        return text.strip()

    def _extract_json(raw: str):
        raw = _strip_fences(raw)
        raw = _truncate_after_json(raw)
        return extract_first_json(raw)

    def _is_garbage_summary(s: str) -> bool:
        if not s or not s.strip():
            return True
        low = s.strip().lower()
        # common phi2/dolphin garbage
        bad = {
            "meeting transcript summary",
            "meeting notes",
            "meeting transcript",
            "meeting summary",
            "summary not provided",
            "meeting transcript not provided",
        }
        if low in bad:
            return True
        # overly short boilerplate
        if len(low) < 12 and "meeting" in low:
            return True
        return False

    def _dedupe_keep_order(lst):
        seen = set()
        out = []
        for x in lst:
            if not isinstance(x, str):
                continue
            tx = " ".join(x.split()).strip()
            if not tx:
                continue
            key = tx.lower()
            if key in seen:
                continue
            seen.add(key)
            out.append(tx)
        return out

    def _take_top_keywords(keywords, n=5):
        """
        Keep top N keywords; prefer heuristic keywords if model keywords are junk/empty.
        """
        kw = [k for k in keywords if isinstance(k, str)]
        kw = [k.strip() for k in kw if k.strip()]
        kw = _dedupe_keep_order(kw)
        return kw[:n]

    # -----------------------------
    # Phase 0: heuristic baseline
    # -----------------------------
    # We'll collect bullets + items, then do one final "structuring" pass.
    all_bullets = []
    all_todo = []
    all_issues = []
    all_suggestions = []
    all_decisions = []
    all_solved = []
    all_keywords = []
    
    combined = {
        "summary": [],
        "bullets": [],
        "keywords": [],
        "todo": [],
        "solved": [],
        "issues": [],
        "suggestions": [],
        "decisions": [],
    }

    # Always seed from heuristics to avoid empties
    for chunk in chunks:
        sents = cached_sent_tokenize(chunk)
        # bullets: first 2 meaningful sentences as baseline bullets
        if sents:
            base = [s.strip() for s in sents[:3] if s and len(s.strip()) > 10]
            for b in base:
                all_bullets.append(b)

        ht, hi, hs = heuristic_extract_items(chunk)
        all_todo.extend(ht)
        all_issues.extend(hi)
        all_suggestions.extend(hs)

    # add heuristic keywords (from your existing extractor)
    try:
        all_keywords.extend(extract_keywords(clean_text, language, n=8))
    except Exception:
        pass

    # -----------------------------
    # Phase 1: chunk-level bullet + candidate item extraction (JSON)
    # -----------------------------
    SYSTEM_PROMPT_CHUNK = (
        "You extract information from meeting transcript text. "
        "Return ONLY valid JSON. No commentary."
    )

    JSON_SCHEMA_CHUNK = (
        "{\n"
        '  "bullets": ["..."],\n'
        '  "todo": ["..."],\n'
        '  "issues": ["..."],\n'
        '  "suggestions": ["..."],\n'
        '  "decisions": ["..."],\n'
        '  "solved": ["..."],\n'
        '  "keywords": ["..."]\n'
        "}\n"
    )

    for idx, chunk in enumerate(chunks, start=1):
        USER_PROMPT_CHUNK = (
            f"LANGUAGE: {language}\n\n"
            "Task:\n"
            "- Extract 3-7 factual bullets.\n"
            "- Extract explicit todo/issues/suggestions/decisions/solved ONLY if clearly stated.\n"
            "- Extract 3-6 keywords (single words or short phrases).\n"
            "- Do NOT invent items. If none exist, use empty lists.\n\n"
            "Return ONLY JSON with this exact structure:\n"
            f"{JSON_SCHEMA_CHUNK}\n"
            "TRANSCRIPT CHUNK:\n"
            f'"""{chunk}"""\n'
        )

        # raw = _call_llm(SYSTEM_PROMPT_CHUNK, USER_PROMPT_CHUNK, max_tokens=550)
        raw = _call_qwen(SYSTEM_PROMPT_CHUNK, USER_PROMPT_CHUNK, max_tokens=550)
        _append_raw(f"CHUNK {idx} RAW", raw)
        
        json_objs = extract_all_json_objects(raw)

        if json_objs:
            merge_structured_content(json_objs, combined)
        else:
            salvage_from_text(raw, combined)

        if not isinstance(combined, dict):
            continue

        # Merge parsed content
        bullets = combined.get("bullets", [])
        todo = combined.get("todo", [])
        issues = combined.get("issues", [])
        suggestions = combined.get("suggestions", [])
        decisions = combined.get("decisions", [])
        solved = combined.get("solved", [])
        keywords = combined.get("keywords", [])

        if isinstance(bullets, list):
            for b in bullets:
                if isinstance(b, str) and not _is_garbage_summary(b) and len(b.strip()) > 10:
                    all_bullets.append(b)

        if isinstance(todo, list):
            all_todo.extend([t for t in todo if isinstance(t, str)])
        if isinstance(issues, list):
            all_issues.extend([t for t in issues if isinstance(t, str)])
        if isinstance(suggestions, list):
            all_suggestions.extend([t for t in suggestions if isinstance(t, str)])
        if isinstance(decisions, list):
            all_decisions.extend([t for t in decisions if isinstance(t, str)])
        if isinstance(solved, list):
            all_solved.extend([t for t in solved if isinstance(t, str)])
        if isinstance(keywords, list):
            all_keywords.extend([k for k in keywords if isinstance(k, str)])

    # Cleanup merged collections
    all_bullets = _dedupe_keep_order(all_bullets)
    all_todo = _dedupe_keep_order(all_todo)
    all_issues = _dedupe_keep_order(all_issues)
    all_suggestions = _dedupe_keep_order(all_suggestions)
    all_decisions = _dedupe_keep_order(all_decisions)
    all_solved = _dedupe_keep_order(all_solved)
    all_keywords = _take_top_keywords(all_keywords, n=8)

    # If bullets are still weak, ensure some baseline bullets exist
    if not all_bullets:
        sents = cached_sent_tokenize(clean_text)
        all_bullets = [s.strip() for s in sents[:5] if s and len(s.strip()) > 10]
        
    def strip_names(text: str) -> str:
        return re.sub(r"\b[A-Z][a-z]+(?: [A-Z][a-z]+){0,2}\b:", "", text).strip()

    all_bullets = [strip_names(b) for b in all_bullets]

    # -----------------------------
    # Phase 2: final structuring pass (ONE call)
    # -----------------------------
    SYSTEM_PROMPT_FINAL = (
        "You are an assistant that writes clean meeting notes from bullet notes. "
        "Use ONLY the provided bullets/items. Do not invent facts. "
        "Return ONLY valid JSON."
    )

    JSON_SCHEMA_FINAL = (
        "{\n"
        '  "summary": "",\n'
        '  "keywords": [],\n'
        '  "todo": [],\n'
        '  "solved": [],\n'
        '  "issues": [],\n'
        '  "suggestions": [],\n'
        '  "decisions": []\n'
        "}\n"
    )
    def filter_bullets(bullets):
        cleaned = []
        for b in bullets:
            low = b.lower()

            if any(x in low for x in [
                "can you", "did you", "are you",
                "screen", "call", "access",
                "calendar", "meeting", "hello"
            ]):
                continue

            cleaned.append(b)

        return cleaned

    all_bullets = filter_bullets(all_bullets)

    # Keep prompts short; phi2/dolphin collapses with long instructions.
    bullets_block = "\n".join(f"- {b}" for b in all_bullets[:60])

    USER_PROMPT_FINAL = (
        f"LANGUAGE: {language}\n\n"
        "You are given extracted meeting bullets and candidate items.\n"
        "Rules:\n"
        "- Summary: write a concise technical meeting summary: 4-6 sentences maximum.\n"
        "   - Do NOT mention people, names, or questions.\n"
        "   - Do NOT describe conversation flow.\n"
        "   - Describe ONLY:\n"
        "   - Databricks setup\n"
        "   - Pipelines\n"
        "   - Config files\n"
        "   - Data flow (bronze/silver/gold)\n"
        "   - Tooling (Azure DevOps, Terraform)\n"
        "   - If something is unclear, omit it. \n"
        "- keywords: 3-6 items.\n"
        "- todo/issues/suggestions/decisions/solved: keep only explicit items.\n"
        "- If a list has no items, return an empty list (do NOT write 'None').\n\n"
        "Return ONLY JSON with this exact structure:\n"
        f"{JSON_SCHEMA_FINAL}\n"
        "BULLETS:\n"
        f"{bullets_block}\n\n"
        "CANDIDATE_TODO:\n"
        + "\n".join(f"- {t}" for t in all_todo[:40]) + "\n\n"
        "CANDIDATE_ISSUES:\n"
        + "\n".join(f"- {t}" for t in all_issues[:40]) + "\n\n"
        "CANDIDATE_SUGGESTIONS:\n"
        + "\n".join(f"- {t}" for t in all_suggestions[:40]) + "\n\n"
        "CANDIDATE_DECISIONS:\n"
        + "\n".join(f"- {t}" for t in all_decisions[:40]) + "\n\n"
        "CANDIDATE_SOLVED:\n"
        + "\n".join(f"- {t}" for t in all_solved[:40]) + "\n"
    )

    # raw_final = _call_llm(SYSTEM_PROMPT_FINAL, USER_PROMPT_FINAL, max_tokens=650)
    raw_final = _call_qwen(SYSTEM_PROMPT_FINAL, USER_PROMPT_FINAL, max_tokens=650)
    _append_raw("FINAL_STRUCTURING_RAW", raw_final)

    json_objs = extract_all_json_objects(raw_final)
    parsed_final = json_objs[0] if json_objs else {}

    if not isinstance(parsed_final, dict):
        parsed_final = {}

    # -----------------------------
    # Final assembly + safeguards
    # -----------------------------
    summary = parsed_final.get("summary", "")
    if not isinstance(summary, str) or _is_garbage_summary(summary):
        # fallback: compress bullets into a readable paragraph
        summary = " ".join(all_bullets[:8])

    final = {
        "summary": clean_summary(summary),
        "keywords": _take_top_keywords(parsed_final.get("keywords", []) or all_keywords, n=6),
        "todo": _dedupe_keep_order(parsed_final.get("todo", []) or all_todo),
        "solved": _dedupe_keep_order(parsed_final.get("solved", []) or all_solved),
        "issues": _dedupe_keep_order(parsed_final.get("issues", []) or all_issues),
        "suggestions": _dedupe_keep_order(parsed_final.get("suggestions", []) or all_suggestions),
        "decisions": _dedupe_keep_order(parsed_final.get("decisions", []) or all_decisions),
    }

    # ensure summary is never empty
    if not final["summary"]:
        sents = cached_sent_tokenize(clean_text)
        final["summary"] = " ".join(s.strip() for s in sents[:3] if s.strip())

    _append_raw("FINAL_EXTRACTED", json.dumps(final, ensure_ascii=False, indent=2))

    return final
